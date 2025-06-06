import os
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
from docx import Document
import re
import markdown
import subprocess
from pathlib import Path
import shutil
from datetime import datetime
import numpy as np
import cv2
import logging
import json
from typing import Dict, List
import sys

VERSION = "1.3.2"  # Major.Minor.Patch - 修复模板处理和日志系统

class MathPaperConverter:
    def __init__(self):
        self.input_dir = Path("input")
        self.output_dir = Path("output")
        self.templates_dir = Path("templates")
        self.temp_dir = Path("temp")
        self.logs_dir = Path("output/logs")
        
        # 设置版本号
        self.version = VERSION
        
        # 调试模式
        self.debug = False
        
        # 确保目录存在
        self.output_dir.mkdir(exist_ok=True)
        self.temp_dir.mkdir(exist_ok=True)
        self.logs_dir.mkdir(exist_ok=True)
        
        # 设置日志
        self._setup_logging()
        
        # 设置Tesseract路径（如果需要）
        if not pytesseract.get_tesseract_version():
            pytesseract.pytesseract.tesseract_cmd = '/usr/local/bin/tesseract'
        
        # 数学符号白名单
        self.math_symbols = "αβγθπ∑∫∏√∈∉∩∪∠⊥∞≈≠±×÷≤≥∂∇∮∴∵∝∞∟∥∦∧∨∩∪⊕⊗⊥∠∡∢∣∤∥∦∧∨∩∪⊕⊗"
        
        # OCR配置
        self.ocr_config = {
            'normal': '--oem 1 --psm 3 -l chi_sim+eng -c preserve_interword_spaces=1 ' + 
                     '-c tessedit_char_blacklist=|~`@#$%^&* ' +
                     '-c textord_tabfind_find_tables=0 ' +
                     '-c tessedit_do_invert=0 ' +
                     '-c textord_heavy_nr=0 ' +
                     '-c textord_min_linesize=2.5 ' +
                     '-c tessedit_char_whitelist="，。！？：；""''（）《》【】、" ' +  # Chinese punctuation
                     '-c tessedit_enable_dict_correction=1 ' +  # Enable dictionary correction
                     '-c language_model_penalty_non_dict_word=0.5 ' +  # Reduce penalty for non-dictionary words
                     '-c language_model_penalty_non_freq_dict_word=0.5',  # Reduce penalty for non-frequent dictionary words
            'math': '--oem 1 --psm 6 -l chi_sim+eng ' +
                   '-c tessedit_char_whitelist="0123456789ABCDEFGHIJKLMNOPQRSTUVWXYZabcdefghijklmnopqrstuvwxyz' +
                   '+-=×÷<>≤≥≠≈∈∉⊆⊂⊇⊃∪∩∅∞∫∑∏√∂∇∆∵∴⊥∥∠∟∣∤∥∦∧∨∩∪⊕⊗→←↑↓↔⇒⇐⇔∀∃∄∎□■△▲○●⊙⊗⊕∘' +
                   '，。！？：；""''（）《》【】、αβγδεζηθικλμνξοπρστυφχψω" ' +  # Added Greek letters and Chinese punctuation
                   '-c preserve_interword_spaces=1 ' +
                   '-c tessedit_char_blacklist=|~`@#$%^&* ' +
                   '-c textord_tabfind_find_tables=0 ' +
                   '-c tessedit_do_invert=0 ' +
                   '-c textord_heavy_nr=0 ' +
                   '-c textord_min_linesize=2.5'
        }
        
        self.logger.info(f"OCR配置已更新：")
        self.logger.info(f"- 普通文本：{self.ocr_config['normal']}")
        self.logger.info(f"- 数学公式：{self.ocr_config['math']}")
        
        # 数学公式校正规则
        self.math_corrections = {
            # Greek letters
            r"a1pha": "α", r"alpha": "α",
            r"beta": "β", r"Beta": "β",
            r"gamma": "γ", r"Gamma": "Γ",
            r"theta": "θ", r"Theta": "Θ",
            r"pi": "π", r"Pi": "Π",
            r"sigma": "σ", r"Sigma": "Σ",
            r"omega": "ω", r"Omega": "Ω",
            
            # Operators and relations
            r"\s?\\int\s?": "∫",
            r"l im": "lim",
            r"inf ty": "∞",
            r"infinity": "∞",
            r"sqrt": "√",
            r"sum": "∑",
            r"prod": "∏",
            r"in": "∈",
            r"notin": "∉",
            r"cap": "∩",
            r"cup": "∪",
            r"angle": "∠",
            r"perp": "⊥",
            r"parallel": "∥",
            r"approx": "≈",
            r"neq": "≠",
            r"leq": "≤",
            r"geq": "≥",
            r"partial": "∂",
            r"nabla": "∇",
            r"oint": "∮",
            r"times": "×",
            r"div": "÷",
            r"pm": "±",
            
            # Common Chinese math terms
            r"（": "(",
            r"）": ")",
            r"［": "[",
            r"］": "]",
            r"｛": "{",
            r"｝": "}",
            r"，": ",",
            r"．": ".",
            r"：": ":",
            r"；": ";",
            r""": "\"",
            r""": "\"",
            r"、": ",",
            r"∵": "因为",
            r"∴": "所以",
            
            # Common formatting fixes
            r"\s*=\s*": " = ",
            r"\s*\+\s*": " + ",
            r"\s*-\s*": " - ",
            r"\s*\*\s*": " × ",
            r"\s*/\s*": " ÷ ",
            r"\(\s+": "(",
            r"\s+\)": ")",
            r"\[\s+": "[",
            r"\s+\]": "]",
            r"\{\s+": "{",
            r"\s+\}": "}"
        }
        
        self.logger.info(f"MathPaperConverter 版本 {VERSION} 初始化完成")
        
    def _setup_logging(self):
        """设置日志系统"""
        self.logger = logging.getLogger('MathPaperConverter')
        self.logger.setLevel(logging.DEBUG)
        
        # 创建一个格式化器
        formatter = logging.Formatter(
            '%(asctime)s - MathPaperConverter %(version)s - %(levelname)s - %(message)s',
            datefmt='%Y-%m-%d %H:%M:%S'
        )
        
        # 添加版本信息到日志记录
        old_factory = logging.getLogRecordFactory()
        def record_factory(*args, **kwargs):
            record = old_factory(*args, **kwargs)
            record.version = self.version
            return record
        logging.setLogRecordFactory(record_factory)
        
        # 生成日志文件名（包含日期和时间）
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        log_filename = f"converter_{timestamp}.log"
        
        # 按日期创建子目录
        date_subdir = datetime.now().strftime('%Y%m%d')
        log_dir = self.logs_dir / date_subdir
        log_dir.mkdir(exist_ok=True)
        
        # 文件处理器
        log_file = log_dir / log_filename
        file_handler = logging.FileHandler(log_file, encoding='utf-8')
        file_handler.setFormatter(formatter)
        file_handler.setLevel(logging.DEBUG)
        self.logger.addHandler(file_handler)
        
        # 控制台处理器
        console_handler = logging.StreamHandler()
        console_handler.setFormatter(formatter)
        console_handler.setLevel(logging.DEBUG)
        self.logger.addHandler(console_handler)
        
        # 记录日志文件位置和版本信息
        self.logger.info(f"日志文件位置: {log_file}")
        self.logger.info(f"当前版本: v{VERSION}")
        
        # 记录系统信息
        self.logger.info(f"操作系统: {os.uname().sysname} {os.uname().release}")
        self.logger.info(f"Python版本: {sys.version}")
        self.logger.info(f"Tesseract版本: {pytesseract.get_tesseract_version()}")
        self.logger.info(f"OpenCV版本: {cv2.__version__}")
        
        # 记录配置信息
        self.logger.info("初始化配置:")
        self.logger.info(f"- 输入目录: {self.input_dir}")
        self.logger.info(f"- 输出目录: {self.output_dir}")
        self.logger.info(f"- 模板目录: {self.templates_dir}")
        self.logger.info(f"- 临时目录: {self.temp_dir}")
        self.logger.info(f"- 日志目录: {self.logs_dir}")

    def _generate_output_filename(self, input_file: Path, suffix: str) -> Path:
        """生成输出文件名，包含时间戳"""
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        return self.output_dir / f"{input_file.stem}_{timestamp}{suffix}"

    def _create_conversion_log(self, input_file: Path, output_files: Dict[str, Path], stats: Dict) -> None:
        """创建转换日志文件"""
        try:
            log_data = {
                "version": self.version,
                "timestamp": datetime.now().isoformat(),
                "input_file": str(input_file),
                "output_files": {k: str(v) for k, v in output_files.items()},
                "stats": stats,
                "optimization_features": {
                    "v1.3.1": [
                        "修复LaTeX模板处理问题",
                        "改进日志系统版本记录",
                        "优化错误处理和报告",
                        "修复编码问题",
                        "添加版本信息到输出文件"
                    ]
                },
                "system_info": {
                    "os": f"{os.uname().sysname} {os.uname().release}",
                    "python_version": sys.version.split()[0],
                    "tesseract_version": str(pytesseract.get_tesseract_version()),
                    "opencv_version": cv2.__version__,
                    "timestamp": datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                },
                "processing_config": {
                    "math_recognition": {
                        "whitelist_symbols": self.math_symbols,
                        "region_detection": True,
                        "auto_correction": True
                    },
                    "image_processing": {
                        "adaptive_threshold": True,
                        "denoising": True,
                        "dpi": 300
                    },
                    "output_formats": ["markdown", "latex", "pdf"]
                }
            }
            
            log_file = self._generate_output_filename(input_file, "_conversion.json")
            with open(log_file, 'w', encoding='utf-8') as f:
                json.dump(log_data, f, ensure_ascii=False, indent=2)
            
            self.logger.info(f"转换日志已保存到：{log_file}")
            
            # 记录当前版本的主要特性
            current_version = f"v{self.version}"
            if current_version in log_data["optimization_features"]:
                self.logger.info(f"当前版本（{current_version}）主要特性：")
                for feature in log_data["optimization_features"][current_version]:
                    self.logger.info(f"- {feature}")
        except Exception as e:
            self.logger.error(f"创建转换日志时出错：{str(e)}")

    def _fix_math_text(self, text: str) -> str:
        """修正数学文本中的常见错误"""
        self.logger.debug("开始修正数学文本")
        original_text = text
        try:
            for pattern, replacement in self.math_corrections.items():
                text = re.sub(pattern, replacement, text)
                if text != original_text:
                    self.logger.debug(f"替换模式 '{pattern}' -> '{replacement}'")
                    original_text = text
            
            self.logger.debug("数学文本修正完成")
            return text
            
        except Exception as e:
            self.logger.error(f"数学文本修正失败: {str(e)}")
            raise

    def _enhance_image(self, image):
        """增强图像质量"""
        # 转换为灰度图
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image
            
        # 自适应直方图均衡化
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(gray)
        
        # 降噪
        denoised = cv2.fastNlMeansDenoising(enhanced)
        
        # 自适应二值化
        binary = cv2.adaptiveThreshold(
            denoised,
            255,
            cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
            cv2.THRESH_BINARY,
            15,  # 块大小
            8    # 常数项
        )
        
        # 形态学操作
        kernel = np.ones((2,2), np.uint8)
        morph = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
        
        return morph
        
    def _detect_math_regions(self, image):
        """检测数学公式区域"""
        try:
            # 转换为灰度图（如果需要）
            if len(image.shape) == 3:
                gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            else:
                gray = image

            # 使用Otsu's二值化
            _, binary = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY_INV + cv2.THRESH_OTSU)

            # 形态学操作以连接相近的数学符号
            kernel = np.ones((3,3), np.uint8)
            dilated = cv2.dilate(binary, kernel, iterations=2)
            
            # 查找轮廓
            contours, _ = cv2.findContours(dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # 过滤和合并数学公式区域
            math_regions = []
            min_area = 100  # 最小区域面积
            
            for contour in contours:
                area = cv2.contourArea(contour)
                if area > min_area:
                    x, y, w, h = cv2.boundingRect(contour)
                    aspect_ratio = w / h
                    
                    # 数学公式通常具有特定的宽高比
                    if 0.2 < aspect_ratio < 10:
                        # 扩展区域以包含完整的公式
                        x = max(0, x - 5)
                        y = max(0, y - 5)
                        w = min(image.shape[1] - x, w + 10)
                        h = min(image.shape[0] - y, h + 10)
                        
                        # 检查区域是否包含数学符号
                        roi = gray[y:y+h, x:x+w]
                        if self._is_math_region(roi):
                            math_regions.append({
                                'x': x, 'y': y, 'w': w, 'h': h,
                                'area': area,
                                'aspect_ratio': aspect_ratio
                            })
            
            # 合并重叠的区域
            merged_regions = []
            while math_regions:
                base = math_regions.pop(0)
                i = 0
                while i < len(math_regions):
                    if self._regions_overlap(base, math_regions[i]):
                        base = self._merge_regions(base, math_regions.pop(i))
                    else:
                        i += 1
                merged_regions.append(base)
            
            # 按y坐标排序
            merged_regions.sort(key=lambda r: r['y'])
            
            return merged_regions
            
        except Exception as e:
            self.logger.error(f"数学公式区域检测失败: {str(e)}")
            self.logger.exception("详细错误信息:")
            return []

    def _is_math_region(self, roi):
        """判断区域是否包含数学符号"""
        try:
            # 对ROI进行预处理
            _, binary = cv2.threshold(roi, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # 计算像素密度
            density = np.sum(binary == 0) / (roi.shape[0] * roi.shape[1])
            
            # 使用OCR进行识别
            text = pytesseract.image_to_string(
                roi,
                config=self.ocr_config['math']
            )
            
            # 检查是否包含数学符号
            math_symbols = set("+-=×÷<>≤≥≠≈∈∉⊆⊂⊇⊃∪∩∅∞∫∑∏√∂∇∆∵∴⊥∥")
            has_math_symbols = any(c in math_symbols for c in text)
            
            # 检查是否包含数字
            has_numbers = any(c.isdigit() for c in text)
            
            # 根据多个条件判断
            return (
                (0.05 < density < 0.5) and  # 适当的像素密度
                (has_math_symbols or has_numbers)  # 包含数学符号或数字
            )
            
        except Exception as e:
            self.logger.error(f"数学区域判断失败: {str(e)}")
            return False

    def _regions_overlap(self, region1, region2):
        """检查两个区域是否重叠"""
        x1, y1, w1, h1 = region1
        x2, y2, w2, h2 = region2
        
        return not (x1 + w1 < x2 or
                   x2 + w2 < x1 or
                   y1 + h1 < y2 or
                   y2 + h2 < y1)
                   
    def _merge_regions(self, region1, region2):
        """合并两个重叠的区域"""
        x1, y1, w1, h1 = region1
        x2, y2, w2, h2 = region2
        
        x = min(x1, x2)
        y = min(y1, y2)
        w = max(x1 + w1, x2 + w2) - x
        h = max(y1 + h1, y2 + h2) - y
        
        return (x, y, w, h)
        
    def _process_math_text(self, text):
        """处理数学文本"""
        # 替换常见的错误识别
        replacements = {
            'pi': 'π',
            'in': '∈',
            'sum': '∑',
            'int': '∫',
            'prod': '∏',
            'sqrt': '√',
            'inf': '∞',
            'approx': '≈',
            'neq': '≠',
            'pm': '±',
            'times': '×',
            'div': '÷',
            'leq': '≤',
            'geq': '≥',
            'partial': '∂',
            'nabla': '∇',
            'therefore': '∴',
            'because': '∵',
            'propto': '∝',
            'perp': '⊥',
            'parallel': '∥',
            # 添加更多数学符号替换
            'alpha': 'α',
            'beta': 'β',
            'gamma': 'γ',
            'delta': 'δ',
            'epsilon': 'ε',
            'theta': 'θ',
            'lambda': 'λ',
            'mu': 'μ',
            'sigma': 'σ',
            'phi': 'φ',
            'omega': 'ω',
            # 添加更多LaTeX命令
            '\\frac': '\\frac',
            '\\sqrt': '\\sqrt',
            '\\sum': '\\sum',
            '\\int': '\\int',
            '\\prod': '\\prod',
            '\\lim': '\\lim',
            '\\infty': '\\infty',
            '\\partial': '\\partial',
            '\\nabla': '\\nabla',
            '\\Delta': '\\Delta',
            '\\alpha': '\\alpha',
            '\\beta': '\\beta',
            '\\gamma': '\\gamma',
            '\\theta': '\\theta',
            '\\lambda': '\\lambda',
            '\\mu': '\\mu',
            '\\sigma': '\\sigma',
            '\\phi': '\\phi',
            '\\omega': '\\omega'
        }
        
        # 替换数学符号
        for old, new in replacements.items():
            text = text.replace(old, new)
            
        # 处理分数
        text = re.sub(r'(\d+)/(\d+)', r'\\frac{\1}{\2}', text)
        text = re.sub(r'\((\d+)\)/\((\d+)\)', r'\\frac{\1}{\2}', text)
        
        # 处理上标
        text = re.sub(r'(\d+)\^(\d+)', r'{\1}^{\2}', text)
        text = re.sub(r'(\w)\^(\w|\d)', r'{\1}^{\2}', text)
        
        # 处理下标
        text = re.sub(r'(\w+)_(\w+)', r'{\1}_{\2}', text)
        text = re.sub(r'(\w)_(\w|\d)', r'{\1}_{\2}', text)
        
        # 处理根号
        text = re.sub(r'sqrt\((.*?)\)', r'\\sqrt{\1}', text)
        text = re.sub(r'√\s*(\d+)', r'\\sqrt{\1}', text)
        text = re.sub(r'√\s*\((.*?)\)', r'\\sqrt{\1}', text)
        
        # 处理特殊数学结构
        text = re.sub(r'lim\s*_\{(.*?)\}', r'\\lim_{\1}', text)
        text = re.sub(r'sum\s*_\{(.*?)\}\^\{(.*?)\}', r'\\sum_{\1}^{\2}', text)
        text = re.sub(r'int\s*_\{(.*?)\}\^\{(.*?)\}', r'\\int_{\1}^{\2}', text)
        
        # 处理括号
        text = re.sub(r'\(', r'\\left(', text)
        text = re.sub(r'\)', r'\\right)', text)
        text = re.sub(r'\[', r'\\left[', text)
        text = re.sub(r'\]', r'\\right]', text)
        
        # 处理省略号
        text = re.sub(r'\.{3,}', r'\\cdots', text)
        
        # 清理多余的空白字符
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        return text
        
    def _clean_text(self, text):
        """清理和规范化文本"""
        # 标点符号替换字典
        punctuation_map = {
            '。': '.',
            '，': ',',
            '、': ',',
            '；': ';',
            '：': ':',
            '"': '"',
            '"': '"',
            ''': "'",
            ''': "'",
            '！': '!',
            '？': '?',
            '（': '(',
            '）': ')',
            '【': '[',
            '】': ']',
            '《': '<',
            '》': '>',
            '—': '-',
            '－': '-',
            '～': '~',
            '〈': '<',
            '〉': '>',
            '「': '[',
            '」': ']',
            '『': '[',
            '』': ']',
            '〔': '(',
            '〕': ')',
            '［': '[',
            '］': ']',
            '｛': '{',
            '｝': '}'
        }
        
        # 数学符号替换字典
        math_symbol_map = {
            '×': '\\times',
            '÷': '\\div',
            '±': '\\pm',
            '∓': '\\mp',
            '≠': '\\neq',
            '≈': '\\approx',
            '≡': '\\equiv',
            '≤': '\\leq',
            '≥': '\\geq',
            '∈': '\\in',
            '∉': '\\notin',
            '⊆': '\\subseteq',
            '⊂': '\\subset',
            '⊇': '\\supseteq',
            '⊃': '\\supset',
            '∪': '\\cup',
            '∩': '\\cap',
            '∅': '\\emptyset',
            '∞': '\\infty',
            '∫': '\\int',
            '∑': '\\sum',
            '∏': '\\prod',
            '√': '\\sqrt',
            '∂': '\\partial',
            '∇': '\\nabla',
            '∆': '\\Delta',
            '∵': '\\because',
            '∴': '\\therefore',
            '⊥': '\\perp',
            '∥': '\\parallel',
            '∠': '\\angle',
            '∟': '\\angle',
            '∣': '\\mid',
            '∤': '\\nmid',
            '∧': '\\wedge',
            '∨': '\\vee',
            '⊕': '\\oplus',
            '⊗': '\\otimes',
            '→': '\\rightarrow',
            '←': '\\leftarrow',
            '↑': '\\uparrow',
            '↓': '\\downarrow',
            '↔': '\\leftrightarrow',
            '⇒': '\\Rightarrow',
            '⇐': '\\Leftarrow',
            '⇔': '\\Leftrightarrow',
            '∀': '\\forall',
            '∃': '\\exists',
            '∄': '\\nexists',
            '∎': '\\blacksquare',
            '□': '\\square',
            '■': '\\blacksquare',
            '△': '\\triangle',
            '▲': '\\blacktriangle',
            '○': '\\circ',
            '●': '\\bullet',
            '⊙': '\\odot',
            '⊗': '\\otimes',
            '⊕': '\\oplus',
            '∘': '\\circ'
        }
        
        # 常见错误替换
        common_errors = {
            '口': '0',
            '〇': '0',
            'O': '0',
            'o': '0',
            'l': '1',
            'I': '1',
            '丨': '1',
            '亅': '1',
            '丿': '1',
            '乙': '2',
            '了': '2',
            '二': '2',
            '三': '3',
            '山': '3',
            '己': '3',
            '已': '3',
            '巳': '3',
            '四': '4',
            '五': '5',
            '六': '6',
            '七': '7',
            '八': '8',
            '九': '9',
            '十': '10',
            '百': '100',
            '千': '1000',
            '万': '10000',
            '亿': '100000000',
            '白': '自',
            '己': '已',
            '已': '己',
            '末': '未',
            '未': '末',
            '木': '本',
            '目': '日',
            '曰': '日',
            '又': '叉',
            '叉': '又',
            '人': '入',
            '入': '人',
            '景': '置',
            '拂': '择',
            '日': '目',
            '马': '号',
            '稽': '稿',
            '冲': '对',
            '诫': '试',
            '拌': '择',
            '札': '本',
            '丝': '卡'
        }
        
        # 替换标点符号
        for cn, en in punctuation_map.items():
            text = text.replace(cn, en)
            
        # 替换数学符号
        for symbol, latex in math_symbol_map.items():
            text = text.replace(symbol, latex)
            
        # 替换常见错误
        for error, correct in common_errors.items():
            text = text.replace(error, correct)
            
        # 规范化空白字符
        text = re.sub(r'\s+', ' ', text)
        text = text.strip()
        
        # 规范化数学公式
        text = re.sub(r'(?<!\\)\$\s+', '$', text)  # 删除$后的空格
        text = re.sub(r'\s+\$', '$', text)  # 删除$前的空格
        
        # 规范化括号
        text = re.sub(r'\(\s+', '(', text)  # 删除(后的空格
        text = re.sub(r'\s+\)', ')', text)  # 删除)前的空格
        text = re.sub(r'\[\s+', '[', text)  # 删除[后的空格
        text = re.sub(r'\s+\]', ']', text)  # 删除]前的空格
        text = re.sub(r'\{\s+', '{', text)  # 删除{后的空格
        text = re.sub(r'\s+\}', '}', text)  # 删除}前的空格
        
        # 规范化数学环境
        text = re.sub(r'\\begin\{([^}]+)\}\s+', r'\\begin{\1}', text)
        text = re.sub(r'\s+\\end\{([^}]+)\}', r'\\end{\1}', text)
        
        # 规范化中文标点
        text = re.sub(r'([，。！？；：、])\s+', r'\1', text)
        
        return text
        
    def _process_math_formulas(self, text):
        """处理数学公式"""
        # 识别数学公式的模式
        math_patterns = [
            (r'\$.*?\$', lambda m: m.group()),  # 行内公式
            (r'\$\$.*?\$\$', lambda m: m.group()),  # 行间公式
            (r'\\begin\{equation\}.*?\\end\{equation\}', lambda m: m.group()),  # equation环境
            (r'\\[.*?\\]', lambda m: m.group()),  # 行间公式
            (r'\\(.*?\\)', lambda m: m.group()),  # 行内公式
        ]
        
        for pattern, handler in math_patterns:
            text = re.sub(pattern, handler, text, flags=re.DOTALL)
            
        return text
        
    def _process_document_structure(self, text):
        """处理文档结构"""
        # 分割成行
        lines = text.split('\n')
        processed_lines = []
        
        for line in lines:
            # 处理标题
            if re.match(r'^第[一二三四五六七八九十]+部分', line):
                line = f'## {line}'
            elif re.match(r'^\d+[、.]', line):
                line = f'### {line}'
                
            # 处理列表
            elif re.match(r'^[A-D][、.．]', line):
                line = f'- {line}'
                
            processed_lines.append(line)
            
        return '\n'.join(processed_lines)

    def process_image(self, image_path_or_image):
        """处理图片文件或图片对象"""
        if isinstance(image_path_or_image, (str, Path)):
            self.logger.info(f"处理图片：{image_path_or_image}")
            # 读取图片
            image = cv2.imread(str(image_path_or_image))
            if image is None:
                raise ValueError(f"无法读取图片：{image_path_or_image}")
        else:
            # 直接使用传入的图片对象
            image = image_path_or_image
            
        try:
            # 图像预处理
            processed_image = self._preprocess_image(image)
            
            # 第一步：使用普通OCR识别所有文本
            text = pytesseract.image_to_string(
                processed_image,
                config='--oem 1 --psm 3 -l chi_sim+eng',
                output_type=pytesseract.Output.DICT
            )
            
            # 获取每个文本块的位置信息
            boxes = pytesseract.image_to_boxes(
                processed_image,
                config='--oem 1 --psm 3 -l chi_sim+eng'
            )
            
            # 第二步：检测数学公式区域
            math_regions = self._detect_math_regions(processed_image)
            
            # 创建一个掩码，标记所有数学公式区域
            mask = np.ones_like(processed_image) * 255
            for x, y, w, h in math_regions:
                cv2.rectangle(mask, (x, y), (x+w, y+h), 0, -1)
            
            # 第三步：分别处理数学公式区域
            math_formulas = []
            for x, y, w, h in math_regions:
                # 提取数学公式区域
                formula_roi = processed_image[y:y+h, x:x+w]
                
                # 增强数学公式区域
                formula_roi = self._enhance_math_region(formula_roi)
                
                # 使用数学公式OCR配置
                formula_text = pytesseract.image_to_string(
                    formula_roi,
                    config='--oem 1 --psm 6 -l chi_sim+eng',
                    output_type=pytesseract.Output.DICT
                )
                
                # 处理数学公式文本
                formula = self._process_math_text(formula_text['text'])
                if formula.strip():
                    math_formulas.append({
                        'text': formula,
                        'position': (x, y, w, h)
                    })
            
            # 第四步：合并结果
            # 将文本按行分组
            lines = []
            current_line = []
            last_y = None
            line_height = 30  # 假设行高为30像素
            
            # 处理每个文本块
            for i, box in enumerate(boxes.split('\n')):
                if not box.strip():
                    continue
                    
                # 解析文本块信息
                parts = box.strip().split()
                if len(parts) < 5:
                    continue
                    
                char = parts[0]
                x = int(parts[1])
                y = int(parts[2])
                w = int(parts[3]) - x
                h = int(parts[4]) - y
                
                # 如果y坐标相差超过行高，认为是新的一行
                if last_y is None or abs(y - last_y) > line_height:
                    if current_line:
                        lines.append(current_line)
                    current_line = []
                    last_y = y
                
                current_line.append({
                    'text': char,
                    'position': (x, y, w, h)
                })
            
            # 添加最后一行
            if current_line:
                lines.append(current_line)
            
            # 将数学公式插入到合适的位置
            result_lines = []
            for line in lines:
                # 计算当前行的y坐标范围
                line_y = sum(block['position'][1] for block in line) / len(line)
                
                # 检查是否有数学公式应该插入在这一行
                formulas_in_line = []
                remaining_formulas = []
                for formula in math_formulas:
                    formula_y = formula['position'][1]
                    if abs(formula_y - line_y) < line_height:
                        formulas_in_line.append(formula)
                    else:
                        remaining_formulas.append(formula)
                
                # 如果这一行有数学公式
                if formulas_in_line:
                    # 按x坐标排序
                    formulas_in_line.sort(key=lambda f: f['position'][0])
                    line_text = ''
                    current_x = 0
                    
                    # 合并文本和公式
                    for block in line:
                        x = block['position'][0]
                        # 检查是否应该先插入公式
                        while formulas_in_line and formulas_in_line[0]['position'][0] < x:
                            formula = formulas_in_line.pop(0)
                            line_text += f" $${formula['text']}$$ "
                        line_text += block['text']
                        current_x = x + block['position'][2]
                    
                    # 添加剩余的公式
                    for formula in formulas_in_line:
                        line_text += f" $${formula['text']}$$ "
                    
                    result_lines.append(line_text)
                else:
                    # 如果这一行没有数学公式，直接合并文本
                    result_lines.append(''.join(block['text'] for block in line))
                
                math_formulas = remaining_formulas
            
            # 处理剩余的数学公式
            for formula in math_formulas:
                result_lines.append(f"$${formula['text']}$$")
            
            # 合并所有行
            text = '\n'.join(result_lines)
            
            # 清理和格式化文本
            text = self._clean_text(text)
            
            # 格式化为Markdown
            final_text = self._format_markdown(text)
            
            return final_text
            
        except Exception as e:
            self.logger.error(f"图片处理失败：{str(e)}")
            raise
            
    def _preprocess_image(self, image):
        """预处理图像"""
        try:
            # 颜色空间标准化
            if len(image.shape) == 3:
                # 确保使用BGR颜色空间
                if image.shape[2] == 4:  # RGBA
                    image = cv2.cvtColor(image, cv2.COLOR_RGBA2BGR)
                elif image.shape[2] == 3:  # 假设是BGR，但可能是RGB
                    # 检查是否需要转换
                    is_rgb = np.mean(image[:,:,0]) > np.mean(image[:,:,2])  # 简单启发式检查
                    if is_rgb:
                        image = cv2.cvtColor(image, cv2.COLOR_RGB2BGR)
            
            # 转换为灰度图
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
            
            # 检查DPI并调整
            dpi = self._check_image_dpi(gray)
            if dpi and dpi < 300:
                scale = 300 / dpi
                gray = cv2.resize(gray, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
            
            # 对比度限制自适应直方图均衡化 (CLAHE)
            # 使用更大的网格尺寸以更好地处理中文字符
            clahe = cv2.createCLAHE(clipLimit=2.5, tileGridSize=(32,32))
            enhanced = clahe.apply(gray)
            
            # 双边滤波以保持边缘锐利度
            denoised = cv2.bilateralFilter(enhanced, d=9, sigmaColor=75, sigmaSpace=75)
            
            # 锐化
            kernel = np.array([[-1,-1,-1],
                             [-1, 9,-1],
                             [-1,-1,-1]])
            sharpened = cv2.filter2D(denoised, -1, kernel)
            
            # Otsu's二值化
            _, binary = cv2.threshold(sharpened, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
            
            # 形态学操作
            # 使用更小的核以更好地保持中文字符的细节
            kernel = np.ones((2,2), np.uint8)
            morph = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
            
            # 去除小噪点但保持文字完整性
            kernel = np.ones((2,2), np.uint8)
            cleaned = cv2.morphologyEx(morph, cv2.MORPH_OPEN, kernel)
            
            # 边缘增强
            edges = cv2.Canny(cleaned, 50, 150)
            enhanced_binary = cv2.addWeighted(cleaned, 0.95, edges, 0.05, 0)
            
            # 保存中间结果（用于调试）
            if self.debug:
                debug_dir = Path('temp/debug')
                debug_dir.mkdir(exist_ok=True)
                timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
                cv2.imwrite(str(debug_dir / f'{timestamp}_01_original.png'), image)
                cv2.imwrite(str(debug_dir / f'{timestamp}_02_gray.png'), gray)
                cv2.imwrite(str(debug_dir / f'{timestamp}_03_enhanced.png'), enhanced)
                cv2.imwrite(str(debug_dir / f'{timestamp}_04_denoised.png'), denoised)
                cv2.imwrite(str(debug_dir / f'{timestamp}_05_sharpened.png'), sharpened)
                cv2.imwrite(str(debug_dir / f'{timestamp}_06_binary.png'), binary)
                cv2.imwrite(str(debug_dir / f'{timestamp}_07_morph.png'), morph)
                cv2.imwrite(str(debug_dir / f'{timestamp}_08_cleaned.png'), cleaned)
                cv2.imwrite(str(debug_dir / f'{timestamp}_09_edges.png'), edges)
                cv2.imwrite(str(debug_dir / f'{timestamp}_10_final.png'), enhanced_binary)
            
            return enhanced_binary
            
        except Exception as e:
            self.logger.error(f"图像预处理失败: {str(e)}")
            self.logger.exception("详细错误信息:")
            return image
        
    def _check_image_dpi(self, image):
        """检查图像DPI"""
        try:
            # 使用PIL获取DPI信息
            pil_image = Image.fromarray(image)
            dpi = pil_image.info.get('dpi', (0, 0))[0]
            return dpi if dpi > 0 else None
        except:
            return None

    def process_pdf(self, pdf_path):
        """处理PDF文件"""
        self.logger.info(f"开始处理PDF文件: {pdf_path}")
        
        try:
            # 使用pdf2image将PDF转换为图片
            images = convert_from_path(str(pdf_path))
            self.logger.info(f"PDF共 {len(images)} 页")
            
            all_text = []
            for i, image in enumerate(images, 1):
                self.logger.info(f"处理第 {i}/{len(images)} 页")
                
                # 将PIL Image转换为OpenCV格式
                cv_image = self._pil_to_cv2(image)
                
                # 处理图片
                text = self.process_image(cv_image)
                all_text.append(text)
                
            # 合并所有页面的文本
            return "\n\n".join(all_text)
            
        except Exception as e:
            self.logger.error(f"PDF处理失败: {str(e)}")
            raise
            
    def _pil_to_cv2(self, pil_image):
        """将PIL Image转换为OpenCV格式"""
        # 转换为RGB格式（如果不是的话）
        if pil_image.mode != 'RGB':
            pil_image = pil_image.convert('RGB')
        
        # 转换为numpy数组
        numpy_image = np.array(pil_image)
        
        # 转换颜色空间从RGB到BGR
        cv_image = cv2.cvtColor(numpy_image, cv2.COLOR_RGB2BGR)
        
        return cv_image

    def process_docx(self, docx_path):
        """处理Word文件，直接提取文本"""
        self.logger.info(f"开始处理Word文件: {docx_path}")
        try:
            doc = Document(docx_path)
            self.logger.info(f"文档共 {len(doc.paragraphs)} 个段落")
            
            text = []
            for i, paragraph in enumerate(doc.paragraphs, 1):
                if paragraph.text.strip():
                    self.logger.debug(f"处理第 {i} 个段落")
                    text.append(paragraph.text)
            
            final_text = '\n'.join(text)
            self.logger.info("Word文件处理完成")
            return self._process_text(final_text)
            
        except Exception as e:
            self.logger.error(f"Word文件处理失败: {str(e)}")
            raise

    def _process_text(self, text):
        """处理提取的文本，识别数学公式和结构"""
        # 清理文本
        text = re.sub(r'\s+', ' ', text)  # 合并多个空格
        text = text.replace('一、', '\n# 一、')  # 添加标题格式
        text = text.replace('二、', '\n# 二、')
        text = text.replace('三、', '\n# 三、')
        text = text.replace('四、', '\n# 四、')
        
        # 处理题号
        text = re.sub(r'(\d+)[.．。]\s*', r'\n### \1. ', text)
        
        # 处理选项
        text = re.sub(r'([A-D])[.．。]\s*', r'\n- \1. ', text)
        
        # 处理数学公式
        text = re.sub(r'([^a-zA-Z])(sin|cos|tan|lim|log|ln|max|min)([^a-zA-Z])', r'\1$\2\3', text)
        text = re.sub(r'([0-9]+)([×÷+-])', r'$\1\2', text)
        text = re.sub(r'([×÷+-])([0-9]+)', r'\1\2$', text)
        
        # 处理特殊数学符号
        for symbol, latex in self.math_corrections.items():
            text = text.replace(symbol, f'${latex}$')
        
        # 清理多余的空行
        text = re.sub(r'\n\s*\n', '\n\n', text)
        
        return text

    def convert_to_markdown(self, content):
        """将处理后的内容转换为Markdown格式"""
        self.logger.info("开始转换为Markdown格式")
        try:
            # 基本文档结构
            markdown_doc = (
                "---\n"
                "title: 数学试卷\n"
                "date: {date}\n"
                "type: math_paper\n"
                "version: {version}\n"
                "---\n\n"
                "{content}"
            )
            
            # 替换内容
            return markdown_doc.format(
                date=datetime.now().strftime("%Y-%m-%d"),
                version=self.version,
                content=content
            )
        except Exception as e:
            self.logger.error(f"Markdown转换失败: {str(e)}")
            raise

    def convert_to_latex(self, content):
        """将处理后的内容转换为LaTeX格式"""
        try:
            # 转义LaTeX特殊字符
            latex_special_chars = {
                '\\': '\\textbackslash{}',
                '{': '\\{',
                '}': '\\}',
                '&': '\\&',
                '#': '\\#',
                '^': '\\^{}',
                '_': '\\_',
                '~': '\\~{}',
                '%': '\\%',
                '$': '\\$',
                '?': '?',
                '!': '!',
                ':': ':',
                ';': ';',
                '<': '$<$',
                '>': '$>$',
                '|': '$|$',
                '*': '$\\ast$'
            }
            
            # 处理内容中的特殊字符
            escaped_content = ''
            i = 0
            in_math_mode = False
            
            while i < len(content):
                if content[i] == '$':
                    in_math_mode = not in_math_mode
                    escaped_content += content[i]
                    i += 1
                elif not in_math_mode:
                    # 检查是否已经是转义序列
                    if content[i:i+2] == '\\\\' or (content[i] == '\\' and i + 1 < len(content) and content[i+1] in latex_special_chars):
                        escaped_content += content[i:i+2]
                        i += 2
                    elif content[i] in latex_special_chars:
                        escaped_content += latex_special_chars[content[i]]
                        i += 1
                    else:
                        # 对于中文字符，不需要特殊处理
                        escaped_content += content[i]
                        i += 1
                else:
                    escaped_content += content[i]
                    i += 1
            
            # 使用内联的LaTeX模板
            latex_template = r'''
\documentclass[12pt,a4paper]{ctexart}
\usepackage{amsmath}
\usepackage{amssymb}
\usepackage{geometry}
\usepackage{graphicx}
\usepackage{enumitem}
\usepackage{fancyhdr}
\usepackage{lastpage}
\usepackage{xeCJK}
\usepackage{unicode-math}

% 设置中文字体
\setCJKmainfont{STSong}
\setCJKsansfont{STHeiti}
\setCJKmonofont{STFangsong}

% 设置页面边距
\geometry{top=2.54cm, bottom=2.54cm, left=3.18cm, right=3.18cm}

% 设置页眉页脚
\pagestyle{fancy}
\fancyhf{}
\renewcommand{\headrulewidth}{0.4pt}
\renewcommand{\footrulewidth}{0.4pt}
\fancyhead[C]{高中数学试卷}
\fancyfoot[C]{第 \thepage 页 / 共 \pageref{LastPage} 页}

% 设置数学公式
\allowdisplaybreaks
\setlength{\parindent}{2em}
\setlength{\parskip}{0.5em}
\linespread{1.5}

% 设置标题和页眉的高度
\setlength{\headheight}{15pt}

% 设置数学符号
\DeclareUnicodeCharacter{2208}{\ensuremath{\in}}
\DeclareUnicodeCharacter{2209}{\ensuremath{\notin}}
\DeclareUnicodeCharacter{2211}{\ensuremath{\sum}}
\DeclareUnicodeCharacter{222B}{\ensuremath{\int}}
\DeclareUnicodeCharacter{220F}{\ensuremath{\prod}}
\DeclareUnicodeCharacter{221A}{\ensuremath{\sqrt}}
\DeclareUnicodeCharacter{221E}{\ensuremath{\infty}}
\DeclareUnicodeCharacter{2248}{\ensuremath{\approx}}
\DeclareUnicodeCharacter{2260}{\ensuremath{\neq}}
\DeclareUnicodeCharacter{2264}{\ensuremath{\leq}}
\DeclareUnicodeCharacter{2265}{\ensuremath{\geq}}
\DeclareUnicodeCharacter{2202}{\ensuremath{\partial}}
\DeclareUnicodeCharacter{2207}{\ensuremath{\nabla}}
\DeclareUnicodeCharacter{2234}{\ensuremath{\therefore}}
\DeclareUnicodeCharacter{2235}{\ensuremath{\because}}

\begin{document}

\begin{center}
    \Large\textbf{高中数学试卷}
\end{center}

\vspace{1em}

%s

\vfill
\begin{center}
    \small{本文档由数学试卷转换器 v%s 自动生成}
\end{center}

\end{document}
'''
            return latex_template % (escaped_content, self.version)
            
        except Exception as e:
            self.logger.error(f"LaTeX转换失败: {str(e)}")
            raise

    def process_file(self, input_file):
        """处理单个输入文件"""
        input_path = self.input_dir / input_file
        self.logger.info(f"开始处理文件：{input_file}")
        
        stats = {"start_time": datetime.now().isoformat()}
        output_files = {}
        
        try:
            # 根据文件类型选择处理方法
            if input_path.suffix.lower() in ['.jpg', '.jpeg', '.png']:
                content = self.process_image(input_path)
                stats["file_type"] = "image"
            elif input_path.suffix.lower() == '.pdf':
                content = self.process_pdf(input_path)
                stats["file_type"] = "pdf"
            elif input_path.suffix.lower() in ['.docx', '.doc']:
                content = self.process_docx(input_path)
                stats["file_type"] = "word"
            else:
                raise ValueError(f"不支持的文件类型: {input_path.suffix}")

            # 输出原始内容以供调试
            self.logger.debug("OCR 识别结果:")
            self.logger.debug("-" * 50)
            self.logger.debug(content)
            self.logger.debug("-" * 50)

            # 生成Markdown文件
            markdown_content = self.convert_to_markdown(content)
            markdown_path = self._generate_output_filename(input_path, ".md")
            
            # 使用UTF-8编码写入文件
            with open(markdown_path, 'w', encoding='utf-8', newline='\n') as f:
                f.write(markdown_content)
            
            output_files["markdown"] = markdown_path
            self.logger.info(f"已生成 Markdown 文件：{markdown_path}")

            # 生成LaTeX文件
            latex_content = self.convert_to_latex(content)
            latex_path = self._generate_output_filename(input_path, ".tex")
            
            with open(latex_path, 'w', encoding='utf-8', newline='\n') as f:
                f.write(latex_content)
                
            output_files["latex"] = latex_path
            self.logger.info(f"已生成 LaTeX 文件：{latex_path}")

            # 生成PDF文件
            pdf_path = latex_path.with_suffix('.pdf')
            self.generate_pdf(latex_content, pdf_path)
            output_files["pdf"] = pdf_path

            stats["end_time"] = datetime.now().isoformat()
            stats["status"] = "success"
            
        except Exception as e:
            self.logger.error(f"处理文件 {input_file} 时出错：{str(e)}")
            stats["end_time"] = datetime.now().isoformat()
            stats["status"] = "error"
            stats["error"] = str(e)
            raise
        
        finally:
            # 创建转换日志
            self._create_conversion_log(input_path, output_files, stats)

    def generate_pdf(self, latex_content, output_path):
        """使用LaTeX生成PDF文件"""
        # 保存 LaTeX 文件到输出目录
        tex_file = output_path.with_suffix('.tex')
        with open(tex_file, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        self.logger.info(f"已生成 LaTeX 文件：{tex_file}")
            
        # 使用临时目录处理编译
        temp_tex_file = self.temp_dir / "temp.tex"
        with open(temp_tex_file, 'w', encoding='utf-8') as f:
            f.write(latex_content)
        
        try:
            # 在临时目录中运行xelatex
            result = subprocess.run(
                ['xelatex', '-interaction=nonstopmode', 'temp.tex'],  # 使用相对路径
                cwd=str(self.temp_dir.absolute()),  # 确保使用绝对路径
                capture_output=True,
                text=True,
                encoding='utf-8'
            )
            
            if result.returncode != 0:
                # 读取详细的日志文件
                log_file = self.temp_dir / "temp.log"
                error_msg = "未知错误"
                if log_file.exists():
                    with open(log_file, 'r', encoding='utf-8') as f:
                        log_content = f.read()
                        # 查找错误信息
                        error_lines = []
                        for line in log_content.split('\n'):
                            if '!' in line or 'Error' in line or 'Fatal' in line:
                                error_lines.append(line)
                        if error_lines:
                            error_msg = '\n'.join(error_lines)
                
                raise RuntimeError(f"LaTeX 编译错误：\n{error_msg}\n\n完整输出：\n{result.stdout}\n{result.stderr}")
            
            # 再次运行以确保交叉引用正确
            result = subprocess.run(
                ['xelatex', '-interaction=nonstopmode', 'temp.tex'],  # 使用相对路径
                cwd=str(self.temp_dir.absolute()),  # 确保使用绝对路径
                capture_output=True,
                text=True,
                encoding='utf-8'
            )
            
            # 移动生成的PDF到目标位置
            temp_pdf = self.temp_dir / "temp.pdf"
            if temp_pdf.exists():
                shutil.move(str(temp_pdf), str(output_path))
                self.logger.info(f"已生成 PDF 文件：{output_path}")
            else:
                raise RuntimeError("PDF 生成失败，请检查 LaTeX 文件中的语法错误")
                
        finally:
            # 清理临时文件，但保留日志文件以供调试
            for temp_file in self.temp_dir.glob("temp.*"):
                if temp_file.suffix != '.log':  # 保留日志文件
                    try:
                        temp_file.unlink()
                    except:
                        pass

    def process_all_files(self):
        """处理输入目录中的所有文件"""
        self.logger.info("开始批量处理文件")
        
        # 检查输入目录是否存在
        if not self.input_dir.exists():
            self.logger.error(f"输入目录 {self.input_dir} 不存在")
            self.input_dir.mkdir(exist_ok=True)
            self.logger.info(f"已创建输入目录 {self.input_dir}")
            self.logger.info("请将需要处理的文件放入输入目录后重新运行程序")
            return
            
        # 列出所有文件
        all_files = list(self.input_dir.iterdir())
        self.logger.info(f"在输入目录中找到 {len(all_files)} 个文件")
        
        found_files = False
        supported_extensions = ['.jpg', '.jpeg', '.png', '.pdf', '.docx', '.doc']
        
        for file_path in all_files:
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                found_files = True
                self.logger.info(f"开始处理文件: {file_path.name}")
                try:
                    self.process_file(file_path.name)
                    self.logger.info(f"成功处理：{file_path.name}")
                except Exception as e:
                    self.logger.error(f"处理文件 {file_path.name} 时出错：{str(e)}")
                    
        if not found_files:
            self.logger.warning(f"在 {self.input_dir} 目录中没有找到支持的文件类型")
            self.logger.info(f"支持的文件类型：{', '.join(supported_extensions)}")
            self.logger.info("请将需要处理的文件放入输入目录后重新运行程序")

    def _enhance_math_region(self, image):
        """增强数学公式区域"""
        # 转换为灰度图
        if len(image.shape) == 3:
            gray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
        else:
            gray = image
            
        # 自适应直方图均衡化
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(gray)
        
        # 锐化
        kernel = np.array([[-1,-1,-1],
                         [-1, 9,-1],
                         [-1,-1,-1]])
        sharpened = cv2.filter2D(enhanced, -1, kernel)
        
        # 二值化
        _, binary = cv2.threshold(sharpened, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
        
        # 形态学操作
        kernel = np.ones((2,2), np.uint8)
        binary = cv2.morphologyEx(binary, cv2.MORPH_CLOSE, kernel)
        
        # 去除小噪点
        kernel = np.ones((3,3), np.uint8)
        binary = cv2.morphologyEx(binary, cv2.MORPH_OPEN, kernel)
        
        return binary
        
    def _find_insert_position(self, lines, y_pos):
        """根据y坐标找到合适的插入位置"""
        # 简单地根据行数估算位置
        estimated_pos = int(y_pos / 50)  # 假设每行大约50像素
        return min(estimated_pos, len(lines))
        
    def _format_markdown(self, content):
        """格式化Markdown文本"""
        # 添加YAML头
        yaml_header = f"""---
title: 数学试卷
date: {datetime.now().strftime('%Y-%m-%d')}
type: math_paper
version: {self.version}
---

"""
        # 处理标题
        lines = content.split('\n')
        formatted_lines = []
        in_list = False
        
        for line in lines:
            # 处理标题
            if re.match(r'^第[一二三四五六七八九十]+部分', line):
                formatted_lines.append(f'\n## {line}\n')
            elif re.match(r'^\d+[、.．]', line):
                formatted_lines.append(f'\n### {line}\n')
            # 处理选项
            elif re.match(r'^[A-D][、.．]', line):
                if not in_list:
                    formatted_lines.append('')
                    in_list = True
                formatted_lines.append(f'- {line}')
            # 处理数学公式
            elif line.startswith('$$'):
                if in_list:
                    formatted_lines.append('')
                    in_list = False
                formatted_lines.append(f'\n{line}\n')
            # 处理普通文本
            else:
                if in_list:
                    formatted_lines.append('')
                    in_list = False
                formatted_lines.append(line)
        
        # 合并行并清理多余的空行
        text = '\n'.join(formatted_lines)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return yaml_header + text.strip()

    def _format_latex(self, content):
        """格式化LaTeX内容"""
        try:
            # 基本文档结构
            latex_doc = (
                "\\documentclass[12pt,a4paper]{ctexart}\n"
                "\\usepackage{amsmath}\n"
                "\\usepackage{amssymb}\n"
                "\\usepackage{geometry}\n"
                "\\usepackage{graphicx}\n"
                "\\usepackage{enumitem}\n"
                "\\usepackage{fancyhdr}\n"
                "\\usepackage{lastpage}\n\n"
                "% 设置页面边距\n"
                "\\geometry{top=2.54cm, bottom=2.54cm, left=3.18cm, right=3.18cm}\n\n"
                "% 设置页眉页脚\n"
                "\\pagestyle{fancy}\n"
                "\\fancyhf{}\n"
                "\\renewcommand{\\headrulewidth}{0.4pt}\n"
                "\\renewcommand{\\footrulewidth}{0.4pt}\n"
                "\\fancyhead[C]{高中数学试卷}\n"
                "\\fancyfoot[C]{第 \\thepage 页 / 共 \\pageref{LastPage} 页}\n\n"
                "% 设置数学公式\n"
                "\\allowdisplaybreaks\n"
                "\\setlength{\\parindent}{2em}\n"
                "\\setlength{\\parskip}{0.5em}\n"
                "\\linespread{1.5}\n\n"
                "% 设置标题和页眉的高度\n"
                "\\setlength{\\headheight}{15pt}\n\n"
                "\\begin{document}\n\n"
                "\\begin{center}\n"
                "    \\Large\\textbf{高中数学试卷}\n"
                "\\end{center}\n\n"
                "\\vspace{1em}\n\n"
                "{content}\n\n"
                "\\end{document}"
            )
            
            # 替换内容
            return latex_doc.format(content=content)
        except Exception as e:
            self.logger.error(f"格式化LaTeX内容时出错：{str(e)}")
            return content

    def _format_log(self, input_file, output_files, stats):
        """格式化日志内容"""
        try:
            log_data = {
                "version": self.version,
                "timestamp": datetime.now().isoformat(),
                "input_file": str(input_file),
                "output_files": {k: str(v) for k, v in output_files.items()},
                "stats": stats,
                "optimization_features": {
                    "v1.3.1": [
                        "修复LaTeX模板处理问题",
                        "改进日志系统版本记录",
                        "优化错误处理和报告",
                        "修复编码问题",
                        "添加版本信息到输出文件"
                    ]
                }
            }
            return json.dumps(log_data, ensure_ascii=False, indent=2)
        except Exception as e:
            self.logger.error(f"格式化日志内容时出错：{str(e)}")
            return "{}"

if __name__ == "__main__":
    try:
        print("数学试卷转换器 v" + VERSION)
        print("正在初始化...")
        converter = MathPaperConverter()
        print(f"请将需要处理的文件放入 {Path('input').absolute()} 目录")
        print("支持的文件类型：PDF、Word文档(.docx/.doc)、图片文件(.jpg/.jpeg/.png)")
        converter.process_all_files()
        print("处理完成，请查看日志文件获取详细信息")
    except Exception as e:
        print(f"程序运行出错：{str(e)}")
        logging.error(f"程序运行出错：{str(e)}")